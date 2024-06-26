import PyPDF2
from django.shortcuts import render
from .forms import PDFUploadForm
from django.http import HttpResponse
from docx import Document
from io import BytesIO
from django.views.decorators.csrf import csrf_exempt
from .prompts import *
import openai
import json

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from django.conf import settings
from datetime import datetime
import os


api_key = "sk-"
client = openai.OpenAI(api_key=api_key)

def generate_content(prompt):

    response = client.chat.completions.create(
        model="gpt-4o",
        temperature=0.00000001,
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content

def extract_text_from_pdf(file):
    reader = PyPDF2.PdfReader(file)
    text = ""
    for page_num in range(len(reader.pages)):
        text += reader.pages[page_num].extract_text()
    return text

def extract_info(resume_text,prompt):
    prompt = f"Given the following resume text: {resume_text}"+prompt
    response_content = generate_content(prompt)
    start_index = response_content.find('{')
    end_index = response_content.rfind('}') + 1
    json_string = response_content[start_index:end_index]
    # Convert the JSON string to a Python dictionary
    extracted_info_json = json.loads(json_string)
    return extracted_info_json

def preprocess_extracted_info(extracted_info):
    def format_item(item):
        if isinstance(item, dict):
            formatted = {}
            for key, value in item.items():
                formatted[key] = format_item(value)
            return formatted
        elif isinstance(item, list):
            return [format_item(subitem) for subitem in item]
        else:
            return item
    
    formatted_info = {}
    for key, value in extracted_info.items():
        if isinstance(value, dict):
            formatted_info[key] = format_item(value)
        elif isinstance(value, list):
            formatted_info[key] = [format_item(item) for item in value]
        else:
            formatted_info[key] = value
    
    return formatted_info

def render_item(item):
    if isinstance(item, str):
        return f"<p>{item}</p>"
    elif isinstance(item, list):
        list_items = ''.join([f"<li class='list-item'>{render_item(subitem)}</li>" for subitem in item])
        return f"<ul>{list_items}</ul>"
    elif isinstance(item, dict):
        dict_items = ''.join([f"<div><span class='subkey'>{subkey}:</span>{render_item(subvalue)}</div>" for subkey, subvalue in item.items()])
        return dict_items
    else:
        return f"<p>{item}</p>"

def add_section_title(doc, title, color=RGBColor(192, 0, 0)):
    section_title = doc.add_paragraph()
    section_title_run = section_title.add_run(title)
    section_title_run.font.size = Pt(18)
    section_title_run.font.color.rgb = color
    section_title_run.font.bold = True
    section_title_run.font.underline = True
    section_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def generate_docx(extracted_info):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run(f"{extracted_info['Trigramme']}\n")
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(192, 0, 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Subtitle (job title)
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"{extracted_info['Métier']}\n")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.color.rgb = RGBColor(255, 0, 0)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Experience with Soft-skills
    experience = doc.add_paragraph()

    # Add Soft-skills if available
    if 'Soft-skills' in extracted_info and extracted_info['Soft-skills']:
        soft_skills_text = ', '.join(extracted_info['Soft-skills'])
        soft_skills_run = experience.add_run(f"{soft_skills_text}\n")
        soft_skills_run.font.size = Pt(12)
        soft_skills_run.font.color.rgb = RGBColor(0, 176, 80)

    # Add years of experience
    experience_value = extracted_info["Années d'expérience"]
    experience_run = experience.add_run(f"{experience_value} ans d'expérience\n")
    experience_run.font.size = Pt(12)
    experience_run.font.color.rgb = RGBColor(0, 176, 80)

    experience.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph()

    add_section_title(doc, 'Savoir-faire clés :')
    savoir_faire_content = doc.add_paragraph()
    for item in extracted_info['Savoir-faire clés']:
        savoir_faire_content_run = savoir_faire_content.add_run(f'• {item}\n')
        savoir_faire_content_run.font.size = Pt(11)

    doc.add_paragraph()

    add_section_title(doc, 'Compétences techniques :')
    competences_table = doc.add_table(rows=1, cols=2)
    competences_table.style = 'Table Grid'

    hdr_cells = competences_table.rows[0].cells
    hdr_cells[0].text = 'Langages de programmation'
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    hdr_cells[1].text = ', '.join(extracted_info['Compétences techniques'].get('Langages de programmation', []))

    for key, value in extracted_info['Compétences techniques'].items():
        if key != 'Langages de programmation' and value:
            row_cells = competences_table.add_row().cells
            row_cells[0].text = key
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[1].text = ', '.join(value)

    if 'Formations' in extracted_info:
        add_section_title(doc, 'Formations :')
        for formation in extracted_info['Formations']:
            if formation['school'] and formation['date'] and formation['title']:
                formation_content = doc.add_paragraph()
                school_run = formation_content.add_run(f"{formation['school']}")
                school_run.font.size = Pt(11)
                school_run.font.bold = True

                date_run = formation_content.add_run(f" ({formation['date']})\n")
                date_run.font.size = Pt(11)
                date_run.font.bold = True
                date_run.font.color.rgb = RGBColor(0, 0, 255)

                title_run = formation_content.add_run(f"{formation['title']}\n")
                title_run.font.size = Pt(11)

    if 'Langues' in extracted_info:
        add_section_title(doc, 'Langues :')
        langues_content = doc.add_paragraph()
        for langue in extracted_info['Langues']:
            if langue['language'] or langue['level']:
                language_run = langues_content.add_run(f"{langue['language']}")
                language_run.font.size = Pt(11)
                language_run.font.bold = True

                level_run = langues_content.add_run(f" ({langue['level']})\n")
                level_run.font.size = Pt(11)
                level_run.font.color.rgb = RGBColor(0, 0, 255)

    if 'Certifications' in extracted_info:
        add_section_title(doc, 'Certifications :')
        certifications_content = doc.add_paragraph()
        for certification in extracted_info['Certifications']:
            if certification['certification_name'] and certification['dates']:
                certification_name_run = certifications_content.add_run(f"{certification['certification_name']}")
                certification_name_run.font.size = Pt(11)
                certification_name_run.font.bold = True

                certification_dates_run = certifications_content.add_run(f" ({certification['dates']})\n")
                certification_dates_run.font.size = Pt(11)
                certification_dates_run.font.bold = True
                certification_dates_run.font.color.rgb = RGBColor(0, 0, 255)

    if 'Expériences professionnelles' in extracted_info:
        add_section_title(doc, 'Expériences professionnelles :')
        for experience in extracted_info['Expériences professionnelles']:
            if experience['job_title'] and experience['client'] and experience['dates']:
                experience_content = doc.add_paragraph()
                job_title_run = experience_content.add_run(f"{experience['job_title']}")
                job_title_run.font.size = Pt(11)
                job_title_run.font.bold = True

                client_dates_run = experience_content.add_run(f" - {experience['client']} ({experience['dates']})\n")
                client_dates_run.font.size = Pt(11)
                client_dates_run.font.bold = True
                client_dates_run.font.color.rgb = RGBColor(0, 0, 255)

                for achievement in experience['achievements']:
                    if achievement:
                        achievement_content_run = experience_content.add_run(f"  • {achievement}\n")
                        achievement_content_run.font.size = Pt(11)

    if 'Projets personnels' in extracted_info:
        add_section_title(doc, 'Projets personnels :')
        for project in extracted_info['Projets personnels']:
            if project['project_title'] and project['context']:
                project_content = doc.add_paragraph()
                project_title_run = project_content.add_run(f"{project['project_title']}")
                project_title_run.font.size = Pt(11)
                project_title_run.font.bold = True

                project_dates_run = project_content.add_run(f" ({project['dates']})\n")
                project_dates_run.font.size = Pt(11)
                project_dates_run.font.bold = True
                project_dates_run.font.color.rgb = RGBColor(0, 0, 255)

                project_context_run = project_content.add_run(f"  • {project['context']}\n")
                project_context_run.font.size = Pt(11)

    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    doc_name = f"{extracted_info['Trigramme']} - {extracted_info['Métier']}.docx"

    return doc_stream, doc_name

# def generate_docx(extracted_info):
#     doc = Document()

#     style = doc.styles['Normal']
#     font = style.font
#     font.name = 'Calibri'
#     font.size = Pt(11)

#     title = doc.add_paragraph()
#     title_run = title.add_run(f"{extracted_info['Trigramme']}\n")
#     title_run.font.size = Pt(24)
#     title_run.font.color.rgb = RGBColor(192, 0, 0)
#     title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#     subtitle = doc.add_paragraph()
#     subtitle_run = subtitle.add_run(f"{extracted_info['Métier']}\n")
#     subtitle_run.font.size = Pt(18)
#     subtitle_run.font.color.rgb = RGBColor(255, 0, 0)
#     subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#     experience = doc.add_paragraph()
#     experience_value = extracted_info["Années d'expérience"]
#     experience_run = experience.add_run(f"{experience_value} ans d'expérience\n")
#     experience_run.font.size = Pt(12)
#     experience_run.font.color.rgb = RGBColor(0, 176, 80)
#     experience.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#     doc.add_paragraph()

#     add_section_title(doc, 'Savoir-faire clés :')
#     savoir_faire_content = doc.add_paragraph()
#     for item in extracted_info['Savoir-faire clés']:
#         savoir_faire_content_run = savoir_faire_content.add_run(f'• {item}\n')
#         savoir_faire_content_run.font.size = Pt(11)

#     doc.add_paragraph()

#     add_section_title(doc, 'Compétences techniques :')
#     competences_table = doc.add_table(rows=1, cols=2)
#     competences_table.style = 'Table Grid'

#     hdr_cells = competences_table.rows[0].cells
#     hdr_cells[0].text = 'Langages de programmation'
#     hdr_cells[0].paragraphs[0].runs[0].font.bold = True
#     hdr_cells[1].text = ', '.join(extracted_info['Compétences techniques'].get('Langages de programmation', []))

#     for key, value in extracted_info['Compétences techniques'].items():
#         if key != 'Langages de programmation' and value:
#             row_cells = competences_table.add_row().cells
#             row_cells[0].text = key
#             row_cells[0].paragraphs[0].runs[0].font.bold = True
#             row_cells[1].text = ', '.join(value)

#     if 'Formations' in extracted_info:
#         add_section_title(doc, 'Formations :')
#         for formation in extracted_info['Formations']:
#             if formation['school'] and formation['date'] and formation['title']:
#                 formation_content = doc.add_paragraph()
#                 school_run = formation_content.add_run(f"{formation['school']}")
#                 school_run.font.size = Pt(11)
#                 school_run.font.bold = True

#                 date_run = formation_content.add_run(f" ({formation['date']})\n")
#                 date_run.font.size = Pt(11)
#                 date_run.font.bold = True
#                 date_run.font.color.rgb = RGBColor(0, 0, 255)

#                 title_run = formation_content.add_run(f"{formation['title']}\n")
#                 title_run.font.size = Pt(11)

#     if 'Langues' in extracted_info:
#         add_section_title(doc, 'Langues :')
#         langues_content = doc.add_paragraph()
#         for langue in extracted_info['Langues']:
#             if langue['language'] or langue['level']:
#                 language_run = langues_content.add_run(f"{langue['language']}")
#                 language_run.font.size = Pt(11)
#                 language_run.font.bold = True

#                 level_run = langues_content.add_run(f" ({langue['level']})\n")
#                 level_run.font.size = Pt(11)
#                 level_run.font.color.rgb = RGBColor(0, 0, 255)

#     if 'Certifications' in extracted_info:
#         add_section_title(doc, 'Certifications :')
#         certifications_content = doc.add_paragraph()
#         for certification in extracted_info['Certifications']:
#             if certification['certification_name'] and certification['dates']:
#                 certification_name_run = certifications_content.add_run(f"{certification['certification_name']}")
#                 certification_name_run.font.size = Pt(11)
#                 certification_name_run.font.bold = True

#                 certification_dates_run = certifications_content.add_run(f" ({certification['dates']})\n")
#                 certification_dates_run.font.size = Pt(11)
#                 certification_dates_run.font.bold = True
#                 certification_dates_run.font.color.rgb = RGBColor(0, 0, 255)

#     if 'Expériences professionnelles' in extracted_info:
#         add_section_title(doc, 'Expériences professionnelles :')
#         for experience in extracted_info['Expériences professionnelles']:
#             if experience['job_title'] and experience['client'] and experience['dates']:
#                 experience_content = doc.add_paragraph()
#                 job_title_run = experience_content.add_run(f"{experience['job_title']}")
#                 job_title_run.font.size = Pt(11)
#                 job_title_run.font.bold = True

#                 client_dates_run = experience_content.add_run(f" - {experience['client']} ({experience['dates']})\n")
#                 client_dates_run.font.size = Pt(11)
#                 client_dates_run.font.bold = True
#                 client_dates_run.font.color.rgb = RGBColor(0, 0, 255)

#                 for achievement in experience['achievements']:
#                     if achievement:
#                         achievement_content_run = experience_content.add_run(f"  • {achievement}\n")
#                         achievement_content_run.font.size = Pt(11)

#     if 'Projets personnels' in extracted_info:
#         add_section_title(doc, 'Projets personnels :')
#         for project in extracted_info['Projets personnels']:
#             if project['project_title'] and project['context']:
#                 project_content = doc.add_paragraph()
#                 project_title_run = project_content.add_run(f"{project['project_title']}")
#                 project_title_run.font.size = Pt(11)
#                 project_title_run.font.bold = True

#                 project_dates_run = project_content.add_run(f" ({project['dates']})\n")
#                 project_dates_run.font.size = Pt(11)
#                 project_dates_run.font.bold = True
#                 project_dates_run.font.color.rgb = RGBColor(0, 0, 255)

#                 project_context_run = project_content.add_run(f"  • {project['context']}\n")
#                 project_context_run.font.size = Pt(11)

#     doc_stream = BytesIO()
#     doc.save(doc_stream)
#     doc_stream.seek(0)

#     doc_name = f"{extracted_info['Trigramme']} - {extracted_info['Métier']}.docx"

#     return doc_stream, doc_name



# def upload_pdf(request):
#     if request.method == 'POST':
#         form = PDFUploadForm(request.POST, request.FILES)
#         if form.is_valid():
#             pdf_file = request.FILES['pdf_file']
#             pdf_text = extract_text_from_pdf(pdf_file)

#             name_info = extract_info(pdf_text, name_prompt)
#             job_info = extract_info(pdf_text, job_prompt)
#             soft_info = extract_info(pdf_text, soft_prompt)
#             xp_info = extract_info(pdf_text, xp_prompt)
#             savoir_info = extract_info(pdf_text, savoir_prompt)
#             hard_info = extract_info(pdf_text, hard_prompt)
            
#             education_info = extract_info(pdf_text, education_prompt)
#             language_info = extract_info(pdf_text, language_prompt)
#             certification_info = extract_info(pdf_text, certification_prompt)
#             professional_info = extract_info(pdf_text, professional_prompt)
#             personal_info = extract_info(pdf_text, personal_project_prompt)

#             extracted_info = {
#                 'Trigramme': name_info['firstName'][0] + name_info['lastName'][-2:],
#                 'Métier': job_info['jobTitle'],
#                 'Soft-skills': soft_info['softSkills'],
#                 "Années d'expérience": xp_info['xp'],
#                 'Savoir-faire clés': savoir_info['SavoirFaireCles'],
#                 'Compétences techniques': hard_info,
#                 'Formations': education_info['formation'],
#                 'Langues': language_info['languages'],
#                 'Certifications': certification_info['certifications'],
#                 'Expériences professionnelles': professional_info['professional_experiences'],
#                 'Projets personnels': personal_info['personal_projects']
#             }

#             processed_info = preprocess_extracted_info(extracted_info)
#             rendered_info = {key: render_item(value) for key, value in processed_info.items()}

#             # Generate DOCX
#             doc_stream, doc_name = generate_docx(extracted_info)
#             request.session['doc_data'] = doc_stream.read().decode('latin1')
#             request.session['doc_name'] = doc_name

#             return render(request, 'extractor/result.html', {'rendered_info': rendered_info})
#     else:
#         form = PDFUploadForm()
#     return render(request, 'extractor/upload.html', {'form': form})

def upload_pdf(request):
    if request.method == 'POST':
        form = PDFUploadForm(request.POST, request.FILES)
        if form.is_valid():
            pdf_file = request.FILES['pdf_file']
            pdf_text = extract_text_from_pdf(pdf_file)

            # Extract information using the prompts
            name_info = extract_info(pdf_text, name_prompt)
            job_info = extract_info(pdf_text, job_prompt)
            soft_info = extract_info(pdf_text, soft_prompt)
            xp_info = extract_info(pdf_text, xp_prompt)
            savoir_info = extract_info(pdf_text, savoir_prompt)
            hard_info = extract_info(pdf_text, hard_prompt)
            education_info = extract_info(pdf_text, education_prompt)
            language_info = extract_info(pdf_text, language_prompt)
            certification_info = extract_info(pdf_text, certification_prompt)
            professional_info = extract_info(pdf_text, professional_prompt)
            personal_info = extract_info(pdf_text, personal_project_prompt)

            # Construct the new dictionary
            extracted_info = {
                'Trigramme': name_info['firstName'][0] + name_info['lastName'][-2:],
                'Métier': job_info['jobTitle'],
                'Soft-skills': soft_info['softSkills'],
                "Années d'expérience": xp_info['xp'],
                'Savoir-faire clés': savoir_info['SavoirFaireCles'],
                'Compétences techniques': hard_info,
                'Formations': education_info['formation'],
                'Langues': language_info['languages'],
                'Certifications': certification_info['certifications'],
                'Expériences professionnelles': professional_info['professional_experiences'],
                'Projets personnels': personal_info['personal_projects']
            }

            processed_info = preprocess_extracted_info(extracted_info)
            rendered_info = {key: render_item(value) for key, value in processed_info.items()}

            # Save the uploaded PDF with the new name
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            pdf_name = f"{name_info['firstName']}_{name_info['lastName']}_{timestamp}.pdf"
            save_path = os.path.join(settings.MEDIA_ROOT, pdf_name)
            with open(save_path, 'wb') as f:
                for chunk in pdf_file.chunks():
                    f.write(chunk)

            # Generate DOCX
            doc_stream, doc_name = generate_docx(extracted_info)
            request.session['doc_data'] = doc_stream.read().decode('latin1')
            request.session['doc_name'] = doc_name

            return render(request, 'extractor/result.html', {'rendered_info': rendered_info})
    else:
        form = PDFUploadForm()
    return render(request, 'extractor/upload.html', {'form': form})

@csrf_exempt
def download_doc(request):
    if request.method == 'POST':
        doc_data = request.session.get('doc_data')
        doc_name = request.session.get('doc_name', 'extracted.docx')
        if doc_data:
            response = HttpResponse(doc_data.encode('latin1'), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{doc_name}"'
            return response
        return HttpResponse("No document available.")
    return HttpResponse("Invalid request.")

